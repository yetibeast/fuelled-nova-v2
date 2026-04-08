"""Tier 2 Valuation Support Document — PwC-style 5-6 page report."""
from __future__ import annotations

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .report_common import (
    NAVY, BLUE, MUTED, WHITE, ORANGE_HEX,
    DISCLAIMER, price_fmt, ref_number, today_str,
    shade, navy_row, alt_shade, font, border_xml, divider,
)

# ── Local helpers ─────────────────────────────────────────────


def _heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        font(r, size=14 if level == 1 else 12, bold=True, color=NAVY)
    return h


def _para(doc, text: str, **kw):
    p = doc.add_paragraph(text)
    for r in p.runs:
        font(r, **kw)
    return p


def _table(doc, headers: list[str], rows: list[tuple]):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                font(r, size=9, bold=True)
    navy_row(t, 0)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    font(r, size=9)
    alt_shade(t)
    return t


def _currency(results: list[dict]) -> str:
    """Extract currency from first result, default CAD."""
    for r in results:
        cur = r.get("structured", {}).get("valuation", {}).get("currency")
        if cur:
            return cur
    return "CAD"


def _fmv_mid(v: dict) -> float:
    """Calculate FMV mid from valuation dict."""
    mid = v.get("fmv_mid")
    if mid:
        return mid
    lo = v.get("fmv_low", 0) or 0
    hi = v.get("fmv_high", 0) or 0
    return (lo + hi) / 2 if (lo or hi) else 0


# ── Main generator ────────────────────────────────────────────


def generate_support_report(
    results: list[dict],
    summary: dict,
    client: str = "",
    synthesis: dict | None = None,
    sections: dict | None = None,
) -> bytes:
    """Generate a PwC-style Valuation Support Document (.docx bytes)."""
    doc = Document()
    ref = ref_number("FVS")
    today = today_str()
    cur = _currency(results)
    pfmt = lambda n: price_fmt(n, cur)

    # Page setup
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Header
    hp = sec.header.paragraphs[0]
    label = f"FUELLED APPRAISALS | FMV Valuation Support | {client}" if client else "FUELLED APPRAISALS | FMV Valuation Support"
    font(hp.add_run(label), size=8, color=MUTED)
    from docx.oxml import parse_xml
    hp._element.get_or_add_pPr().append(parse_xml(border_xml("bottom", sz="8")))

    # Footer
    fp = sec.footer.paragraphs[0]
    footer_text = f"Confidential | Prepared by Fuelled Energy Marketing Inc. | Effective Date: {today}"
    font(fp.add_run(footer_text), size=8, color=RGBColor(0x99, 0x99, 0x99))
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Cover ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(p.add_run("FMV VALUATION SUPPORT"), size=26, bold=True, color=NAVY)
    if client:
        p = doc.add_paragraph()
        font(p.add_run(client), size=16, bold=True, color=BLUE)
    doc.add_paragraph()
    _para(doc, f"Date: {today}")
    _para(doc, f"Reference: {ref}")
    _para(doc, "Prepared By: Fuelled Energy Marketing Inc.")
    divider(doc)

    # ── Valuation Summary Table ──
    total_mid = sum(_fmv_mid(r.get("structured", {}).get("valuation", {})) for r in results)
    # Group by category
    cats: dict[str, dict] = {}
    for r in results:
        v = r.get("structured", {}).get("valuation", {})
        cat = v.get("category") or r.get("title", "Equipment")
        entry = cats.setdefault(cat, {"units": 0, "mid_total": 0})
        entry["units"] += 1
        entry["mid_total"] += _fmv_mid(v)

    rows = []
    for cat, d in cats.items():
        mid_per = d["mid_total"] / d["units"] if d["units"] else 0
        rows.append((cat, str(d["units"]), pfmt(mid_per), pfmt(d["mid_total"])))
    rows.append(("TOTAL", str(len(results)), "", pfmt(total_mid)))
    _table(doc, ["Category", "Units", "FMV Mid/Unit", "FMV Mid Subtotal"], rows)

    # Offer comparison line
    if synthesis and synthesis.get("offer_amount"):
        offer = synthesis["offer_amount"]
        delta = offer - total_mid
        pct = (delta / total_mid * 100) if total_mid else 0
        sign = "+" if delta >= 0 else ""
        doc.add_paragraph()
        _para(doc, f"Buyer's Offer: {pfmt(offer)} | Offer vs. FMV Mid: {sign}{pfmt(delta)} ({sign}{pct:.1f}%)",
              size=10, bold=True)

    doc.add_page_break()

    # ── Section 1: Equipment Identification ──
    _heading(doc, "Section 1: Equipment Identification")
    # Build parameter table from first result (portfolio-level)
    equip_names = ", ".join(r.get("title", "") for r in results[:5])
    if len(results) > 5:
        equip_names += f" (+{len(results) - 5} more)"
    id_rows = [
        ("Equipment", equip_names),
        ("Total Items", str(summary.get("total", len(results)))),
        ("Valuation Date", today),
        ("Reference", ref),
    ]
    _table(doc, ["Parameter", "Details"], id_rows)
    doc.add_paragraph()

    # ── Executive Summary (Claude-enriched) ──
    if sections and sections.get("executive_summary"):
        _heading(doc, "Executive Summary")
        for para_text in sections["executive_summary"].split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                _para(doc, para_text, size=10)
        doc.add_paragraph()

    # ── Section 2: Equipment Description ──
    _heading(doc, "Section 2: Equipment Description")
    if sections and sections.get("equipment_description"):
        ed = sections["equipment_description"]
        # Overview paragraph
        if ed.get("overview"):
            _para(doc, ed["overview"], size=10)
        # Specs table
        specs = ed.get("specs_table") or []
        if specs:
            doc.add_paragraph()
            spec_rows = [(s.get("component", ""), s.get("specification", "")) for s in specs]
            _table(doc, ["Component", "Specification"], spec_rows)
        # Notes
        if ed.get("notes"):
            doc.add_paragraph()
            _para(doc, ed["notes"], size=9, italic=True, color=MUTED)
    else:
        # Fallback: existing sparse condition text
        for r in results:
            v = r.get("structured", {})
            title = r.get("title", "Equipment")
            cond = v.get("condition_assessment", "")
            p = doc.add_paragraph()
            font(p.add_run(title), size=11, bold=True, color=BLUE)
            if cond:
                _para(doc, cond, size=9)
            val = v.get("valuation", {})
            _para(doc, f"FMV Range: {pfmt(val.get('fmv_low'))} \u2014 {pfmt(val.get('fmv_high'))} | "
                  f"Mid: {pfmt(_fmv_mid(val))} | Confidence: {val.get('confidence', r.get('confidence', 'N/A'))}",
                  size=9, color=MUTED)

    # ── Valuation Methodology (Claude-enriched) ──
    if sections and sections.get("valuation_methodology"):
        doc.add_paragraph()
        vm = sections["valuation_methodology"]
        _heading(doc, "Valuation Methodology")
        # Approach paragraph
        if vm.get("approach"):
            _para(doc, vm["approach"], size=10)

        # RCN Derivation
        rcn_d = vm.get("rcn_derivation") or {}
        if rcn_d.get("narrative"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            font(p.add_run("RCN Derivation"), size=11, bold=True, color=NAVY)
            _para(doc, rcn_d["narrative"], size=10)
        if rcn_d.get("components"):
            doc.add_paragraph()
            comp_rows = [(c.get("component", ""), c.get("cost_range", "")) for c in rcn_d["components"]]
            if rcn_d.get("total_rcn"):
                comp_rows.append(("TOTAL RCN (Estimated)", rcn_d["total_rcn"]))
            _table(doc, ["Component", "Cost Range"], comp_rows)
        if rcn_d.get("notes"):
            _para(doc, rcn_d["notes"], size=9, italic=True, color=MUTED)

        # Depreciation
        dep = vm.get("depreciation") or {}
        if dep.get("formula") or dep.get("factors"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            font(p.add_run("Depreciation Analysis"), size=11, bold=True, color=NAVY)
        if dep.get("formula"):
            _para(doc, dep["formula"], size=10, bold=True)
        if dep.get("factors"):
            dep_rows = [
                (f.get("factor", ""),
                 str(f.get("multiplier", "")),
                 f.get("rationale", ""))
                for f in dep["factors"]
            ]
            _table(doc, ["Factor", "Multiplier", "Rationale"], dep_rows)
        if dep.get("notes"):
            _para(doc, dep["notes"], size=9, italic=True, color=MUTED)

    doc.add_page_break()

    # ── Section 3: Comparable Sales Evidence ──
    _heading(doc, "Section 3: Comparable Sales Evidence")

    if sections and sections.get("market_comparables"):
        mc = sections["market_comparables"]
        # Overview paragraph
        if mc.get("overview"):
            _para(doc, mc["overview"], size=10)
            doc.add_paragraph()
        # Use Claude's enriched listings for the table
        listings = mc.get("listings") or []
        if listings:
            comp_rows = [
                (
                    lg.get("source", ""),
                    lg.get("description", lg.get("title", "")),
                    str(lg.get("year", "")),
                    lg.get("location", ""),
                    lg.get("price", "") if isinstance(lg.get("price"), str) else pfmt(lg.get("price")),
                    lg.get("notes", ""),
                )
                for lg in listings
            ]
            _table(doc, ["Source", "Equipment", "Year", "Location", "Price", "Notes"], comp_rows)
        else:
            _para(doc, "No comparable sales data available.", size=9, italic=True, color=MUTED)
        # Analysis paragraphs
        if mc.get("analysis"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            font(p.add_run("Comparable Analysis"), size=11, bold=True, color=NAVY)
            for para_text in mc["analysis"].split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    _para(doc, para_text, size=9)
    else:
        # Fallback: pull from structured data
        comp_rows = []
        for r in results:
            for c in r.get("structured", {}).get("comparables", []):
                comp_rows.append((
                    c.get("source", ""),
                    c.get("title", ""),
                    str(c.get("year", "")),
                    c.get("location", ""),
                    pfmt(c.get("price")),
                    c.get("notes", ""),
                ))
        if comp_rows:
            _table(doc, ["Source", "Equipment", "Year", "Location", "Price", "Notes"], comp_rows)
        else:
            _para(doc, "No comparable sales data available.", size=9, italic=True, color=MUTED)

        # Key Comparable Observations (fallback)
        all_drivers = []
        for r in results:
            all_drivers.extend(r.get("structured", {}).get("key_value_drivers", []))
        if all_drivers:
            doc.add_paragraph()
            p = doc.add_paragraph()
            font(p.add_run("Key Comparable Observations"), size=11, bold=True, color=NAVY)
            for d in all_drivers:
                _para(doc, f"\u2022 {d}", size=9)
    doc.add_page_break()

    # ── Section 4: Offer Analysis & Key Valuation Factors (conditional) ──
    if synthesis and synthesis.get("offer_amount"):
        _heading(doc, "Section 4: Offer Analysis & Key Valuation Factors")
        factors = synthesis.get("key_factors", [])
        for i, f_ in enumerate(factors, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. ")
            font(run, size=10, bold=True)
            run2 = p.add_run(f_)
            font(run2, size=10)
        doc.add_paragraph()

    # ── Section 5: Valuation Reconciliation ──
    _heading(doc, "Section 5: Valuation Reconciliation")
    total_lo = summary.get("total_fmv_low", 0) or 0
    total_hi = summary.get("total_fmv_high", 0) or 0
    recon_rows = [
        ("FMV Low", pfmt(total_lo)),
        ("FMV Mid", pfmt(total_mid)),
        ("FMV High", pfmt(total_hi)),
    ]
    if synthesis and synthesis.get("offer_amount"):
        offer = synthesis["offer_amount"]
        delta = offer - total_mid
        pct = (delta / total_mid * 100) if total_mid else 0
        sign = "+" if delta >= 0 else ""
        recon_rows.append(("Buyer's Offer", pfmt(offer)))
        recon_rows.append(("Offer vs. FMV Mid", f"{sign}{pct:.1f}%"))
    _table(doc, ["Metric", "Value"], recon_rows)
    doc.add_paragraph()

    # Claude-enriched FMV details
    if sections and sections.get("fair_market_value"):
        fmv = sections["fair_market_value"]
        # Scenarios table
        if fmv.get("scenarios"):
            p = doc.add_paragraph()
            font(p.add_run("Valuation Scenarios"), size=11, bold=True, color=NAVY)
            scen_rows = [
                (
                    s.get("scenario", ""),
                    pfmt(s.get("fmv_low")),
                    pfmt(s.get("fmv_high")),
                    str(s.get("confidence", "")),
                )
                for s in fmv["scenarios"]
            ]
            _table(doc, ["Scenario", "FMV Low", "FMV High", "Confidence"], scen_rows)
            doc.add_paragraph()
        # List pricing table
        if fmv.get("list_pricing"):
            lp = fmv["list_pricing"]
            p = doc.add_paragraph()
            font(p.add_run("List Pricing Guidance"), size=11, bold=True, color=NAVY)
            lp_rows = [
                (
                    pfmt(lp.get("fmv_midpoint")),
                    str(lp.get("list_premium", "")),
                    pfmt(lp.get("recommended_list")),
                    pfmt(lp.get("walkaway_floor")),
                ),
            ]
            _table(doc, ["FMV Midpoint", "List Premium", "Recommended List", "Walk-Away"], lp_rows)
            doc.add_paragraph()
        # Overhaul economics
        if fmv.get("overhaul_economics"):
            p = doc.add_paragraph()
            font(p.add_run("Overhaul Economics"), size=11, bold=True, color=NAVY)
            _para(doc, fmv["overhaul_economics"], size=10)

    # ── Section 6: Key Value Drivers ──
    _heading(doc, "Section 6: Key Value Drivers")
    # Collect drivers from structured data (used both in fallback and enriched paths)
    all_drivers: list[str] = []
    for r in results:
        all_drivers.extend(r.get("structured", {}).get("key_value_drivers", []))
    if all_drivers:
        for i, d in enumerate(all_drivers, 1):
            _para(doc, f"{i}. {d}", size=10)
    else:
        _para(doc, "No key value drivers identified.", size=9, italic=True, color=MUTED)
    doc.add_paragraph()

    # ── Assumptions ──
    if sections and sections.get("assumptions"):
        _heading(doc, "Assumptions & Limiting Conditions")
        assumptions = sections["assumptions"]
        for i, a in enumerate(assumptions, 1):
            _para(doc, f"{i}. {a}", size=9)
        doc.add_paragraph()

    # ── Sources ──
    _heading(doc, "Sources")
    if sections and sections.get("sources"):
        for s in sections["sources"]:
            _para(doc, f"\u2022 {s}", size=9)
    else:
        all_sources: list[str] = []
        for r in results:
            for s in r.get("structured", {}).get("sources", []):
                if s not in all_sources:
                    all_sources.append(s)
        if all_sources:
            for s in all_sources:
                _para(doc, f"\u2022 {s}", size=9)
        else:
            _para(doc, "No external sources cited.", size=9, italic=True, color=MUTED)
    doc.add_paragraph()

    # ── Disclaimer ──
    divider(doc)
    _para(doc, DISCLAIMER, size=8, italic=True, color=MUTED)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
