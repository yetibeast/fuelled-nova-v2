"""Professional .docx report generation for equipment valuations."""
from __future__ import annotations
import datetime
import re
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Colour palette (locked — clients like the current format) ───
NAVY = RGBColor(0x1A, 0x1A, 0x2E)
BLUE = RGBColor(0x00, 0x77, 0xB6)
ORANGE_HEX = "E85D04"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_HEX = "F5F5F5"
MUTED = RGBColor(0x66, 0x66, 0x66)


def _ref():
    d = datetime.date.today()
    return f"FV-{d.year}-{d.month:02d}{d.day:02d}"


def _today():
    return datetime.date.today().strftime("%B %d, %Y")


def _price(n):
    if n is None:
        return "[N/A]"
    return f"${n:,.0f} CAD"


# ── Low-level helpers ───────────────────────────────────────────

def _shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))


def _navy_row(table, idx):
    for cell in table.rows[idx].cells:
        _shade(cell, "1A1A2E")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.bold = True


def _alt_shade(table, start=1):
    for i in range(start, len(table.rows)):
        if i % 2 == 0:  # even data rows (0-indexed from start) get shading
            for cell in table.rows[i].cells:
                _shade(cell, GRAY_HEX)


def _font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color


def _border_xml(tag, color=ORANGE_HEX, sz="12"):
    return (f'<w:pBdr {nsdecls("w")}>'
            f'<w:{tag} w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
            f'</w:pBdr>')


def _divider(doc):
    p = doc.add_paragraph()
    p._element.get_or_add_pPr().append(parse_xml(_border_xml("bottom")))


def _heading(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        _font(r, size=14, bold=True, color=NAVY)


def _subheading(doc, text):
    p = doc.add_paragraph()
    _font(p.add_run(text), size=11, bold=True, color=BLUE)


def _para(doc, text, **kw):
    p = doc.add_paragraph(text)
    for r in p.runs:
        _font(r, **kw)
    return p


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                _font(r, size=9, bold=True)
    _navy_row(t, 0)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    _font(r, size=9)
    _alt_shade(t)
    return t


def _no_border(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.append(tblPr)
    borders = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders))


def _factor_display(f):
    val = f.get("value", "")
    if isinstance(val, (int, float)) and val <= 5:
        return f"{val:.2f}x"
    return str(val)


def _strip_markdown(text: str) -> str:
    """Remove markdown formatting from text for clean Word document insertion."""
    if not text:
        return ""
    # Remove code blocks
    text = re.sub(r'```[\s\S]*?```', '', text)
    # Remove header markers
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Remove bold/italic markers
    text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', text)
    # Remove horizontal rules
    text = re.sub(r'^-{3,}$', '', text, flags=re.MULTILINE)
    # Remove bullet markers at start of lines
    text = re.sub(r'^\s*[\*\-]\s+', '', text, flags=re.MULTILINE)
    # Collapse multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ── Main entry point ───────────────────────────────────────────

def generate_report(structured: dict, response_text: str, user_message: str) -> bytes:
    """Generate a professional .docx valuation report. Returns file bytes."""
    doc = Document()
    ref, today = _ref(), _today()

    # ── Page setup ─────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # Page header with orange bottom border
    hp = sec.header.paragraphs[0]
    _font(hp.add_run(f"Fuelled Energy Marketing | Equipment Valuation Report | {ref}"),
          size=8, color=MUTED)
    hp._element.get_or_add_pPr().append(parse_xml(_border_xml("bottom", sz="8")))

    # Page footer
    fp = sec.footer.paragraphs[0]
    _font(fp.add_run("Confidential | fuelled.com"),
          size=8, color=RGBColor(0x99, 0x99, 0x99))
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Extract structured data ────────────────────────
    v = structured.get("valuation", {})
    comps = structured.get("comparables", [])
    risks = structured.get("risks", [])
    factors = v.get("factors", [])
    fmv_low, fmv_high, rcn = v.get("fmv_low"), v.get("fmv_high"), v.get("rcn")
    confidence = v.get("confidence", "MEDIUM")
    equip_title = (
        v.get("title")
        or v.get("type")
        or "Equipment Valuation"
    )

    # ═══════════════════════════════════════════════════
    # 1. COVER PAGE
    # ═══════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run("EQUIPMENT VALUATION REPORT"), size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    _font(p.add_run(equip_title), size=14, bold=True, color=BLUE)
    _divider(doc)
    doc.add_paragraph()

    info = doc.add_table(rows=5, cols=2)
    for i, (label, val) in enumerate([
        ("Prepared For:", "[Client Name]"),
        ("Contact:", "[Contact]"),
        ("Prepared By:", "Fuelled Energy Marketing Inc."),
        ("Date:", today),
        ("Reference Number:", ref),
    ]):
        info.rows[i].cells[0].text = label
        info.rows[i].cells[1].text = val
        for p in info.rows[i].cells[0].paragraphs:
            for r in p.runs:
                _font(r, bold=True)
        for p in info.rows[i].cells[1].paragraphs:
            for r in p.runs:
                _font(r)
    _no_border(info)

    doc.add_paragraph()
    doc.add_paragraph()
    _para(doc, "CONFIDENTIAL \u2014 This document contains proprietary valuation methodology "
          "and market data. Intended for authorized recipients only.",
          size=9, italic=True, color=MUTED)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 2. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════
    _heading(doc, "EXECUTIVE SUMMARY")
    cond_display = next(
        (str(f["value"]) for f in factors if "cond" in f.get("label", "").lower()),
        "B (Good) \u2014 Assumed")
    _table(doc, ["Valuation Summary", ""], [
        ("Replacement Cost New (RCN)", _price(rcn)),
        ("Fair Market Value Range",
         f"{_price(fmv_low)} \u2014 {_price(fmv_high)}" if fmv_low else "[Not provided]"),
        ("Valuation Basis", "Cost Approach (RCN \u00d7 Depreciation) + Market Comparables"),
        ("Condition Assessment", cond_display),
        ("Effective Date", today),
    ])
    doc.add_paragraph()
    # Build a clean executive summary from structured data
    summary_parts = [
        f"This report provides a fair market value assessment for {equip_title}."
    ]
    if fmv_low and fmv_high:
        summary_parts.append(
            f"Based on RCN methodology with market validation against {len(comps)} comparable "
            f"listing{'s' if len(comps) != 1 else ''}, the estimated FMV range "
            f"is {_price(fmv_low)} \u2014 {_price(fmv_high)} with {confidence} confidence."
        )
    elif comps:
        summary_parts.append(
            f"{len(comps)} market comparable(s) were identified to support the valuation."
        )
    _para(doc, " ".join(summary_parts))
    _divider(doc)

    # ═══════════════════════════════════════════════════
    # 3. EQUIPMENT DESCRIPTION
    # ═══════════════════════════════════════════════════
    _heading(doc, "EQUIPMENT DESCRIPTION")
    _subheading(doc, "Technical Specifications")
    specs = [("Equipment", equip_title)]
    if rcn:
        specs.append(("Replacement Cost New", _price(rcn)))
    if fmv_low:
        specs.append(("Estimated FMV Range", f"{_price(fmv_low)} \u2014 {_price(fmv_high)}"))
    for f in factors:
        specs.append((f.get("label", "Factor"), _factor_display(f)))
    specs.append(("Confidence Level", confidence))
    _table(doc, ["Specification", "Value"], specs)
    _divider(doc)

    # ═══════════════════════════════════════════════════
    # 4. VALUATION METHODOLOGY — RCN vs Market Comp
    # ═══════════════════════════════════════════════════
    _heading(doc, "VALUATION METHODOLOGY")

    # RCN Approach
    _subheading(doc, "Cost Approach (RCN)")
    if rcn:
        _para(doc, f"The Replacement Cost New (RCN) of {_price(rcn)} represents the estimated "
              "cost to replace this equipment with a functionally equivalent new unit at current "
              "market prices in Canadian dollars.")
    else:
        _para(doc, "RCN data not available for this equipment type.")

    _subheading(doc, "Depreciation Adjustments")
    if factors:
        _table(doc, ["Factor", "Multiplier", "Rationale"],
               [(f.get("label", ""), _factor_display(f),
                 _strip_markdown(f.get("rationale", "Applied per Fuelled depreciation curves")))
                for f in factors])
    else:
        _para(doc, "Depreciation factors not available for this valuation.")
    doc.add_paragraph()

    if fmv_low and rcn:
        mid = (fmv_low + fmv_high) / 2
        p = doc.add_paragraph()
        _font(p.add_run(
            f"FMV = RCN ({_price(rcn)}) \u00d7 Combined Factor ({mid / rcn:.4f}) = {_price(mid)}"),
            size=10, bold=True, color=BLUE)

    # Market Comp Approach (side by side comparison)
    _subheading(doc, "Market Comparable Approach")
    if comps:
        asking_prices = [c.get("price", 0) for c in comps if c.get("price")]
        if asking_prices:
            avg_asking = sum(asking_prices) / len(asking_prices)
            adj_low = avg_asking * 0.80
            adj_high = avg_asking * 0.90
            _table(doc, ["Metric", "Cost Approach (RCN)", "Market Comp Approach"], [
                ("Basis", "RCN × Depreciation", f"{len(comps)} comparable listings"),
                ("Value Range",
                 f"{_price(fmv_low)} — {_price(fmv_high)}" if fmv_low else "[N/A]",
                 f"{_price(adj_low)} — {_price(adj_high)}"),
                ("Adjustments", "Age, condition, hours, service", "10-20% asking price discount"),
                ("Data Source", "Fuelled RCN tables", "Fuelled marketplace (36,000+ listings)"),
            ])
        else:
            _para(doc, "Market comparables found but pricing data insufficient for approach comparison.")
    else:
        _para(doc, "No direct market comparables available for approach comparison.")
    _divider(doc)

    # ═══════════════════════════════════════════════════
    # 5. MARKET COMPARABLES
    # ═══════════════════════════════════════════════════
    _heading(doc, "MARKET COMPARABLES")
    if comps:
        _table(doc, ["Description", "Price", "Year", "Location", "Source", "URL"],
               [(_strip_markdown(c.get("title", "?")), _price(c.get("price")),
                 str(c.get("year", "-")), c.get("location", "-"),
                 c.get("source", "Fuelled"), c.get("url", "-")) for c in comps])
        doc.add_paragraph()
        _subheading(doc, "Comparables Analysis")
        _para(doc, f"{len(comps)} comparable(s) identified from the Fuelled marketplace database "
              "of 36,000+ listings. Note: listed asking prices typically reflect a 10\u201320% "
              "premium over actual transaction values.")
    else:
        _para(doc, "No direct market comparables identified. Common for high-specification or "
              "custom-engineered packages that typically trade through direct negotiation.")

    # Sources section
    _subheading(doc, "Data Sources")
    source_set = set()
    for c in comps:
        url = c.get("url", "")
        src = c.get("source", "Fuelled")
        if url:
            source_set.add(f"{src}: {url}")
        elif src:
            source_set.add(src)
    if source_set:
        for s in source_set:
            _para(doc, f"\u2022 {s}", size=9)
    else:
        _para(doc, "Sources: Fuelled marketplace database", size=9)
    _divider(doc)

    # ═══════════════════════════════════════════════════
    # 6. FAIR MARKET VALUE
    # ═══════════════════════════════════════════════════
    _heading(doc, "FAIR MARKET VALUE")
    if fmv_low and fmv_high:
        mid = (fmv_low + fmv_high) / 2
        _table(doc, ["Scenario", "FMV Low", "FMV High", "Confidence"],
               [("Primary Scenario", _price(fmv_low), _price(fmv_high), confidence)])
        doc.add_paragraph()
        _subheading(doc, "Recommended List Pricing")
        list_p, walkaway = mid * 1.12, fmv_low * 0.92
        _table(doc, ["Metric", "Value"], [
            ("FMV Midpoint", _price(mid)),
            ("List Premium (12%)", _price(list_p)),
            ("Recommended List Price", _price(list_p)),
            ("Walk-Away Floor", _price(walkaway)),
        ])
    else:
        _para(doc, "Fair market value could not be determined from available data.")
    _divider(doc)

    # ═══════════════════════════════════════════════════
    # 7. KEY ASSUMPTIONS AND LIMITING CONDITIONS
    # ═══════════════════════════════════════════════════
    _heading(doc, "KEY ASSUMPTIONS AND LIMITING CONDITIONS")
    for i, a in enumerate([
        "This valuation is based on information provided and has not been independently "
        "verified through physical inspection.",
        "Equipment condition has been assumed based on reported age, hours, and service "
        "type unless otherwise stated.",
        "All values are expressed in Canadian Dollars (CAD) unless otherwise noted.",
        "Market comparables are based on listed asking prices, which typically exceed "
        "actual transaction values by 10\u201320%.",
        "RCN values are derived from Fuelled\u2019s proprietary reference tables and "
        "industry benchmarks.",
        "Depreciation factors follow Fuelled\u2019s standardized curves based on equipment "
        "class, age, condition, hours, and service type.",
        "This valuation does not account for transportation, installation, commissioning, "
        "or decommissioning costs.",
        "Values assume the equipment is free and clear of liens, encumbrances, or "
        "environmental liabilities.",
    ], 1):
        _para(doc, f"{i}. {a}")
    _divider(doc)

    # ═══════════════════════════════════════════════════
    # 8. DISCLAIMER
    # ═══════════════════════════════════════════════════
    _heading(doc, "DISCLAIMER")
    _para(doc,
          "This document represents Fuelled Energy Marketing Inc.\u2019s opinion of value "
          "based on available market data and proprietary methodology. It is not a certified "
          "appraisal and should not be construed as such. This report is intended solely for "
          "the use of the addressee and should not be relied upon by third parties without the "
          "express written consent of Fuelled Energy Marketing Inc. Market conditions are "
          "subject to change and values presented herein are effective as of the date noted above.",
          size=9, italic=True, color=MUTED)
    _divider(doc)

    # ═══════════════════════════════════════════════════
    # 9. SIGNATURE BLOCK
    # ═══════════════════════════════════════════════════
    doc.add_paragraph()
    _para(doc, "Prepared by:", bold=True)
    doc.add_paragraph("_" * 40)
    for line in [
        "Fuelled Energy Marketing Inc.",
        "Harsh Kansara | harsh.kansara@fuelled.com",
        "Mark Le Dain | mark.ledain@fuelled.com",
        "fuelled.com",
    ]:
        _para(doc, line)

    # ── Serialize ──────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
