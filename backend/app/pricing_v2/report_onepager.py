"""Tier 1 one-pager FMV report generator — single-page valuation summary."""
from __future__ import annotations

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .report_common import (
    NAVY, BLUE, ORANGE_HEX, WHITE, MUTED,
    FOOTER_LINE, price_fmt, today_str,
    shade, navy_row, alt_shade, font, divider,
)


def generate_onepager(
    structured: dict,
    user_message: str,
    client: str = "",
    synthesis: dict | None = None,
) -> bytes:
    """Generate a one-page FMV summary .docx and return its bytes."""
    val = structured.get("valuation", {})
    currency = val.get("currency", "CAD")
    if synthesis:
        currency = synthesis.get("currency", currency)
    date = today_str()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # 1 — Header line
    p = doc.add_paragraph()
    r = p.add_run(f"FUELLED APPRAISALS | FMV Valuation Support | {client}" if client else "FUELLED APPRAISALS | FMV Valuation Support")
    font(r, size=8, color=MUTED)

    divider(doc)

    # 2 — Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FAIR MARKET VALUE")
    font(r, size=20, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VALUATION SUPPORT DOCUMENT")
    font(r, size=11, color=BLUE)

    divider(doc)

    # 3 — Equipment line
    p = doc.add_paragraph()
    r = p.add_run(user_message)
    font(r, size=10, bold=True)

    # 4 — Info line
    info_parts = []
    if client:
        info_parts.append(client)
    info_parts.append(f"Effective Date: {date}")
    info_parts.append(f"All Values in {currency}")
    p = doc.add_paragraph()
    r = p.add_run(" | ".join(info_parts))
    font(r, size=9, color=MUTED)

    # 5 — Valuation Summary table
    rows = _build_rows(structured, val, currency, synthesis)
    cols = ["Category", "Units", "FMV Mid / Unit", "FMV Mid Subtotal"]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"

    # Header
    for i, hdr in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = hdr
    navy_row(table, 0)

    # Data rows
    for ri, row_data in enumerate(rows, start=1):
        for ci, val_text in enumerate(row_data):
            table.rows[ri].cells[ci].text = val_text

    alt_shade(table, start=1)

    # Orange total row (last data row)
    if rows:
        last = len(rows)
        for cell in table.rows[last].cells:
            shade(cell, ORANGE_HEX)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE
                    r.font.bold = True

    doc.add_paragraph()  # spacer

    # 6 — Basis of Value
    p = doc.add_paragraph()
    r = p.add_run("Basis of Value: Fair Market Value, As-Is/Where-Is, Orderly Liquidation (Bulk Sale).")
    font(r, size=9, italic=True, color=MUTED)

    divider(doc)

    # 7 — Footer
    p = doc.add_paragraph()
    r = p.add_run(f"Confidential | Fuelled Energy Marketing Inc. | Effective Date: {date}")
    font(r, size=8, color=MUTED)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_rows(structured: dict, val: dict, currency: str, synthesis: dict | None) -> list[list[str]]:
    """Build table rows: category breakdown if synthesis, else single row + total."""
    rows: list[list[str]] = []

    if synthesis and synthesis.get("category_breakdown"):
        for cat in synthesis["category_breakdown"]:
            rows.append([
                cat.get("category", "Equipment"),
                str(cat.get("count", 1)),
                price_fmt(cat.get("fmv_mid_per_unit"), currency),
                price_fmt(cat.get("fmv_mid_subtotal"), currency),
            ])
        totals = synthesis.get("totals", {})
        rows.append([
            "TOTAL",
            str(totals.get("count", "")),
            "",
            price_fmt(totals.get("fmv_mid"), currency),
        ])
    else:
        fmv_mid = val.get("fmv_mid", 0)
        rows.append(["Equipment", "1", price_fmt(fmv_mid, currency), price_fmt(fmv_mid, currency)])
        rows.append(["TOTAL", "1", "", price_fmt(fmv_mid, currency)])

    return rows
