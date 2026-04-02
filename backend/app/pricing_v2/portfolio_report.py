"""Portfolio report: multi-item .docx for batch valuations."""
from __future__ import annotations
import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x1A, 0x1A, 0x2E)
BLUE = RGBColor(0x00, 0x77, 0xB6)
ORANGE_HEX = "E85D04"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_HEX = "F5F5F5"
MUTED = RGBColor(0x66, 0x66, 0x66)


def _ref():
    d = datetime.date.today()
    return f"FV-PORTFOLIO-{d.year}-{d.month:02d}{d.day:02d}"


def _today():
    return datetime.date.today().strftime("%B %d, %Y")


def _price(n):
    if n is None or n == 0:
        return "[N/A]"
    return f"${n:,.0f} CAD"


def _shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))


def _navy_row(table, idx):
    for cell in table.rows[idx].cells:
        _shade(cell, "1A1A2E")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.bold = True


def _alt_shade(table, start=1):
    for i in range(start, len(table.rows)):
        if i % 2 == 0:
            for cell in table.rows[i].cells:
                _shade(cell, GRAY_HEX)


def _font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color


def _border_xml(tag, color=ORANGE_HEX, sz="12"):
    return (f'<w:pBdr {nsdecls("w")}>'
            f'<w:{tag} w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
            f'</w:pBdr>')


def _heading(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        _font(r, size=14, bold=True, color=NAVY)


def _para(doc, text, **kw):
    p = doc.add_paragraph(text)
    for r in p.runs:
        _font(r, **kw)
    return p


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                _font(r, size=9, bold=True)
    _navy_row(t, 0)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    _font(r, size=9)
    _alt_shade(t)
    return t


def generate_portfolio_report(results: list[dict], batch_summary: dict) -> bytes:
    doc = Document()
    ref, today = _ref(), _today()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    hp = sec.header.paragraphs[0]
    _font(hp.add_run(f"Fuelled Energy Marketing | Portfolio Valuation Report | {ref}"), size=8, color=MUTED)
    hp._element.get_or_add_pPr().append(parse_xml(_border_xml("bottom", sz="8")))
    fp = sec.footer.paragraphs[0]
    _font(fp.add_run("Confidential | fuelled.com"), size=8, color=RGBColor(0x99, 0x99, 0x99))
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_lo = batch_summary.get("total_fmv_low", 0)
    total_hi = batch_summary.get("total_fmv_high", 0)

    # 1. Cover
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run("PORTFOLIO VALUATION REPORT"), size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    _font(p.add_run(f"{batch_summary.get('total', 0)} Equipment Items"), size=14, bold=True, color=BLUE)
    doc.add_paragraph()
    _para(doc, f"Date: {today}")
    _para(doc, f"Reference: {ref}")
    _para(doc, "Prepared By: Fuelled Energy Marketing Inc.")
    doc.add_page_break()

    # 2. Executive Summary
    _heading(doc, "EXECUTIVE SUMMARY")
    _para(doc, f"This portfolio valuation covers {batch_summary.get('total', 0)} equipment items. "
          f"Of these, {batch_summary.get('completed', 0)} were successfully valued and "
          f"{batch_summary.get('failed', 0)} could not be assessed. "
          f"The aggregate FMV range is {_price(total_lo)} \u2014 {_price(total_hi)}.")
    doc.add_paragraph()
    _table(doc, ["Metric", "Value"], [
        ("Total Items", str(batch_summary.get("total", 0))),
        ("Successfully Valued", str(batch_summary.get("completed", 0))),
        ("Failed / Skipped", str(batch_summary.get("failed", 0))),
        ("Total FMV Low", _price(total_lo)),
        ("Total FMV High", _price(total_hi)),
        ("Effective Date", today),
    ])
    doc.add_page_break()

    # 3. Portfolio Summary Table
    _heading(doc, "PORTFOLIO SUMMARY TABLE")
    rows = []
    for i, r in enumerate(results, 1):
        v = r.get("structured", {}).get("valuation", {})
        title = v.get("title") or v.get("type") or r.get("title", f"Item {i}")
        rows.append((str(i), title, _price(v.get("fmv_low")), _price(v.get("fmv_high")),
                      r.get("confidence", "LOW")))
    _table(doc, ["#", "Equipment", "FMV Low", "FMV High", "Confidence"], rows)
    doc.add_page_break()

    # 4. Category Breakdown
    _heading(doc, "CATEGORY BREAKDOWN")
    cats: dict[str, dict] = {}
    for r in results:
        v = r.get("structured", {}).get("valuation", {})
        cat = v.get("category", "Unknown") or "Unknown"
        cats.setdefault(cat, {"count": 0, "lo": 0, "hi": 0})
        cats[cat]["count"] += 1
        cats[cat]["lo"] += v.get("fmv_low", 0) or 0
        cats[cat]["hi"] += v.get("fmv_high", 0) or 0
    cat_rows = [(c, str(d["count"]), _price(d["lo"]), _price(d["hi"])) for c, d in cats.items()]
    _table(doc, ["Category", "Count", "Total FMV Low", "Total FMV High"], cat_rows)
    doc.add_page_break()

    # 5. Individual Line Items
    _heading(doc, "INDIVIDUAL LINE ITEMS")
    for i, r in enumerate(results, 1):
        v = r.get("structured", {}).get("valuation", {})
        title = v.get("title") or r.get("title", f"Item {i}")
        comps = r.get("structured", {}).get("comparables", [])
        rationale = (r.get("response", "") or "")[:200]
        p = doc.add_paragraph()
        _font(p.add_run(f"{i}. {title}"), size=11, bold=True, color=BLUE)
        _para(doc, f"RCN: {_price(v.get('rcn'))} | FMV: {_price(v.get('fmv_low'))} \u2014 "
              f"{_price(v.get('fmv_high'))} | Confidence: {r.get('confidence', 'LOW')} | "
              f"Comps: {len(comps)}", size=9)
        if rationale:
            _para(doc, rationale[:200], size=9, italic=True, color=MUTED)
    doc.add_page_break()

    # 6. Methodology
    _heading(doc, "METHODOLOGY")
    _para(doc, "All valuations use Fuelled's RCN (Replacement Cost New) methodology with "
          "deterministic depreciation factors for age, condition, hours, and service type. "
          "Market validation is performed against 36,000+ active listings.")
    doc.add_paragraph()

    # 7. Assumptions & Disclaimer
    _heading(doc, "ASSUMPTIONS & DISCLAIMER")
    _para(doc, "Equipment condition assumed based on reported data unless verified. "
          "All values in CAD. Listed asking prices typically exceed transaction values by 10-20%. "
          "This is not a certified appraisal.", size=9, italic=True, color=MUTED)
    doc.add_paragraph()
    _para(doc, "Prepared by:", bold=True)
    for line in ["Fuelled Energy Marketing Inc.", "valuations@fuelled.com"]:
        _para(doc, line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
