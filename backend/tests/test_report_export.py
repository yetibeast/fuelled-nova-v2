"""Regression test for generate_report — Shawn's workflow depends on docx export."""
from io import BytesIO

from docx import Document

from app.pricing_v2.report import generate_report


def _sample_structured() -> dict:
    """Shape matches what the pricing agent returns for a valuation query."""
    return {
        "valuation": {
            "type": "Ariel JGK/4 Gas Compressor Package",
            "title": "2015 Waukesha L7042 / Ariel JGK/4 800HP Package",
            "currency": "USD",
            "fmv_low": 165000,
            "fmv_mid": 195000,
            "fmv_high": 223000,
            "rcn": 480000,
            "confidence": "HIGH",
            "list_price": 218000,
            "walkaway": 152000,
            "factors": [
                {"label": "Age (11 years)", "value": 0.65},
                {"label": "Condition B", "value": 0.85},
            ],
        },
        "comparables": [
            {
                "title": "2003 Waukesha L5774LT / Ariel JGK/4",
                "price": 184000,
                "currency": "USD",
                "year": "2003",
                "location": "Grande Prairie, AB",
                "source": "Fuelled.com",
            },
        ],
        "risks": ["18,000 hours — verify major overhaul history"],
        "market_context": "Strong demand for 600-1000HP gas compression.",
        "assumptions": ["Condition B assumed"],
        "sources": ["Fuelled.com"],
    }


def test_generate_report_returns_valid_docx():
    docx_bytes = generate_report(
        structured=_sample_structured(),
        response_text="Narrative goes here.",
        user_message="Valuation request",
    )
    assert len(docx_bytes) > 10_000, "DOCX suspiciously small"

    doc = Document(BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)

    # Key client-facing sections must be present
    assert "EQUIPMENT VALUATION REPORT" in text
    assert "EXECUTIVE SUMMARY" in text
    assert "VALUATION METHODOLOGY" in text
    assert "FAIR MARKET VALUE" in text

    # FMV range must render somewhere (paragraphs or tables)
    all_cell_text = " ".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    combined = text + " " + all_cell_text
    assert "$165,000" in combined and "$223,000" in combined, "FMV range missing"
    assert "$480,000" in combined, "RCN missing"

    # At least one comparable row must reach the comps table
    assert "Waukesha L5774LT" in all_cell_text, "Comparable not rendered"


def test_generate_report_handles_empty_comps():
    """Discovery-style queries may have no comparables; report must still render."""
    structured = _sample_structured()
    structured["comparables"] = []
    docx_bytes = generate_report(
        structured=structured,
        response_text="",
        user_message="Valuation request",
    )
    assert len(docx_bytes) > 10_000
    doc = Document(BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "FAIR MARKET VALUE" in text
