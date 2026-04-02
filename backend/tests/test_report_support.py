"""Tests for Tier 2 valuation support report (PwC format)."""
import pytest
from docx import Document
from io import BytesIO

from app.pricing_v2.report_support import generate_support_report


def _single_result():
    return [{"title": "Gas Driven Zedi Pump Skid", "structured": {
        "valuation": {"fmv_low": 1500, "fmv_high": 2500, "fmv_mid": 2000,
                       "confidence": "HIGH", "currency": "CAD"},
        "comparables": [{"title": "Kudu 5.7L", "price": 2000, "source": "Fuelled",
                          "url": "https://fuelled.com/62549"}],
        "condition_assessment": "Mixed, 60-70% complete",
        "key_value_drivers": ["Identical comps"],
        "sources": ["Fuelled.com #62549"],
        "assumptions": ["Not inspected"],
    }, "confidence": "HIGH"}]


def _summary(n=1):
    return {"total": n, "completed": n, "total_fmv_low": 1500 * n, "total_fmv_high": 2500 * n}


def test_generates_valid_docx():
    result = generate_support_report(_single_result(), _summary(), client="PwC / Longrun")
    assert isinstance(result, bytes)
    assert result[:2] == b'PK'


def test_multiple_items():
    results = [{"title": f"Item {i}", "structured": {
        "valuation": {"fmv_low": 1000 * i, "fmv_high": 2000 * i, "currency": "CAD"}},
        "confidence": "MEDIUM"} for i in range(1, 6)]
    summary = {"total": 5, "completed": 5, "total_fmv_low": 15000, "total_fmv_high": 30000}
    result = generate_support_report(results, summary)
    assert isinstance(result, bytes)


def test_header_contains_client():
    data = generate_support_report(_single_result(), _summary(), client="PwC / Longrun")
    doc = Document(BytesIO(data))
    header_text = doc.sections[0].header.paragraphs[0].text
    assert "PwC / Longrun" in header_text
    assert "FUELLED APPRAISALS" in header_text


def test_footer_contains_confidential():
    data = generate_support_report(_single_result(), _summary())
    doc = Document(BytesIO(data))
    footer_text = doc.sections[0].footer.paragraphs[0].text
    assert "Confidential" in footer_text
    assert "Fuelled Energy Marketing" in footer_text


def test_disclaimer_present():
    data = generate_support_report(_single_result(), _summary())
    doc = Document(BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    full = " ".join(texts)
    assert "opinion of value" in full


def test_offer_analysis_section_with_offer():
    synthesis = {"offer_amount": 50000, "offer_vs_fmv_mid_pct": 12.5,
                 "key_factors": ["Strong demand", "Low supply"]}
    data = generate_support_report(_single_result(), _summary(),
                                    client="Test", synthesis=synthesis)
    doc = Document(BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    full = " ".join(texts)
    assert "Offer Analysis" in full or "offer" in full.lower()


def test_no_offer_section_without_synthesis():
    data = generate_support_report(_single_result(), _summary())
    doc = Document(BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    full = " ".join(texts)
    # Section 4 should not appear
    assert "Offer Analysis" not in full


def test_comps_table_present():
    data = generate_support_report(_single_result(), _summary())
    doc = Document(BytesIO(data))
    tables = doc.tables
    # Should have at least valuation summary + comps tables
    assert len(tables) >= 2


def test_sources_section():
    data = generate_support_report(_single_result(), _summary())
    doc = Document(BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    full = " ".join(texts)
    assert "Fuelled.com #62549" in full
