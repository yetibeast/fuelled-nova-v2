"""Tests for shared report helpers in report_common.py."""
from __future__ import annotations
import datetime
from unittest.mock import patch

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from app.pricing_v2.report_common import (
    NAVY, BLUE, ORANGE_HEX, WHITE, GRAY_HEX, MUTED,
    DISCLAIMER, FOOTER_LINE,
    price_fmt, ref_number, today_str,
    shade, navy_row, alt_shade, font, border_xml, divider,
)


# ── price_fmt ─────────────────────────────────────────────────

def test_price_cad():
    assert price_fmt(150000, "CAD") == "$150,000 CAD"


def test_price_usd():
    assert price_fmt(150000, "USD") == "$150,000 USD"


def test_price_none():
    assert price_fmt(None, "CAD") == "[N/A]"


def test_price_zero():
    assert price_fmt(0, "CAD") == "[N/A]"


def test_price_default():
    assert price_fmt(50000) == "$50,000 CAD"


# ── ref_number ────────────────────────────────────────────────

@patch("app.pricing_v2.report_common.datetime")
def test_ref_number_default(mock_dt):
    mock_dt.date.today.return_value = datetime.date(2026, 4, 1)
    assert ref_number() == "FV-2026-0401"


@patch("app.pricing_v2.report_common.datetime")
def test_ref_number_custom_prefix(mock_dt):
    mock_dt.date.today.return_value = datetime.date(2026, 4, 1)
    assert ref_number("PV") == "PV-2026-0401"


# ── today_str ─────────────────────────────────────────────────

def test_today_str():
    # Just verify it returns a non-empty string in expected format
    result = today_str()
    assert len(result) > 0
    # Should contain a comma (e.g. "April 01, 2026")
    assert "," in result


# ── shade ─────────────────────────────────────────────────────

def test_shade_applies_fill():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade(cell, "FF0000")
    xml = cell._tc.xml
    assert 'w:fill="FF0000"' in xml


# ── navy_row ──────────────────────────────────────────────────

def test_navy_row_applies_shading_and_style():
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    for cell in t.rows[0].cells:
        cell.text = "Header"
    navy_row(t, 0)
    xml = t.rows[0].cells[0]._tc.xml
    assert 'w:fill="1A1A2E"' in xml
    for cell in t.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                assert r.font.color.rgb == WHITE
                assert r.font.bold is True


# ── alt_shade ─────────────────────────────────────────────────

def test_alt_shade_even_rows():
    doc = Document()
    t = doc.add_table(rows=4, cols=1)
    for i in range(4):
        t.rows[i].cells[0].text = f"row {i}"
    alt_shade(t, start=1)
    # Row 2 (even index) should be shaded
    assert 'w:fill="F5F5F5"' in t.rows[2].cells[0]._tc.xml
    # Row 1 (odd index) should NOT be shaded
    assert 'w:fill="F5F5F5"' not in t.rows[1].cells[0]._tc.xml


# ── font ──────────────────────────────────────────────────────

def test_font_sets_properties():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("test")
    font(run, size=12, bold=True, italic=True, color=NAVY)
    assert run.font.name == "Arial"
    assert run.font.size == Pt(12)
    assert run.font.bold is True
    assert run.font.italic is True
    assert run.font.color.rgb == NAVY


def test_font_defaults():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("test")
    font(run)
    assert run.font.name == "Arial"
    assert run.font.size == Pt(10)


# ── border_xml ────────────────────────────────────────────────

def test_border_xml_default():
    xml = border_xml("bottom")
    assert "w:bottom" in xml
    assert ORANGE_HEX in xml
    assert 'w:sz="12"' in xml


def test_border_xml_custom():
    xml = border_xml("top", color="FF0000", sz="8")
    assert "w:top" in xml
    assert "FF0000" in xml
    assert 'w:sz="8"' in xml


# ── divider ───────────────────────────────────────────────────

def test_divider_adds_paragraph():
    doc = Document()
    count_before = len(doc.paragraphs)
    divider(doc)
    assert len(doc.paragraphs) == count_before + 1


# ── constants ─────────────────────────────────────────────────

def test_constants_exist():
    assert isinstance(NAVY, RGBColor)
    assert isinstance(BLUE, RGBColor)
    assert isinstance(ORANGE_HEX, str)
    assert isinstance(WHITE, RGBColor)
    assert isinstance(GRAY_HEX, str)
    assert isinstance(MUTED, RGBColor)
    assert isinstance(DISCLAIMER, str)
    assert "Fuelled" in DISCLAIMER
    assert "valuations@fuelled.com" in FOOTER_LINE
