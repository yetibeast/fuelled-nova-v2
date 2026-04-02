"""Tests for Tier 1 one-pager report generator."""
from __future__ import annotations

import pytest
from docx import Document
from io import BytesIO

from app.pricing_v2.report_onepager import generate_onepager


# ── basic generation ─────────────────────────────────────────

def test_generates_valid_docx():
    data = {"valuation": {"fmv_low": 100000, "fmv_high": 200000, "fmv_mid": 150000, "confidence": "HIGH", "currency": "CAD"}}
    result = generate_onepager(data, user_message="Test equipment", client="Test Client")
    assert isinstance(result, bytes)
    assert result[:2] == b'PK'


def test_usd_currency():
    data = {"valuation": {"fmv_low": 50000, "fmv_high": 80000, "fmv_mid": 65000, "currency": "USD"}}
    result = generate_onepager(data, user_message="US equipment")
    assert isinstance(result, bytes)
    assert result[:2] == b'PK'


# ── content verification ─────────────────────────────────────

def _doc_text(data, **kwargs) -> str:
    """Generate a one-pager and return all paragraph + table text concatenated."""
    buf = generate_onepager(data, **kwargs)
    doc = Document(BytesIO(buf))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_header_contains_client():
    data = {"valuation": {"fmv_mid": 100000, "currency": "CAD"}}
    text = _doc_text(data, user_message="Ariel JGK4 compressor", client="Acme Oil")
    assert "Acme Oil" in text
    assert "FUELLED APPRAISALS" in text


def test_title_block():
    data = {"valuation": {"fmv_mid": 100000, "currency": "CAD"}}
    text = _doc_text(data, user_message="test equip", client="C")
    assert "FAIR MARKET VALUE" in text
    assert "VALUATION SUPPORT DOCUMENT" in text


def test_equipment_description_shown():
    data = {"valuation": {"fmv_mid": 50000, "currency": "CAD"}}
    text = _doc_text(data, user_message="2018 Caterpillar 3516 generator set")
    assert "2018 Caterpillar 3516 generator set" in text


def test_basis_of_value():
    data = {"valuation": {"fmv_mid": 50000, "currency": "CAD"}}
    text = _doc_text(data, user_message="test")
    assert "Fair Market Value" in text
    assert "Orderly Liquidation" in text


def test_footer_present():
    data = {"valuation": {"fmv_mid": 50000, "currency": "CAD"}}
    text = _doc_text(data, user_message="test")
    assert "Confidential" in text
    assert "Fuelled Energy Marketing" in text


def test_currency_in_info_line():
    data = {"valuation": {"fmv_mid": 50000, "currency": "USD"}}
    text = _doc_text(data, user_message="test", client="X")
    assert "USD" in text


# ── batch / synthesis mode ───────────────────────────────────

def test_batch_with_synthesis():
    data = {"results": [
        {"valuation": {"fmv_mid": 100000, "currency": "CAD"}},
        {"valuation": {"fmv_mid": 200000, "currency": "CAD"}},
    ]}
    synthesis = {
        "currency": "CAD",
        "category_breakdown": [
            {"category": "Compressors", "count": 2, "fmv_mid_per_unit": 150000, "fmv_mid_subtotal": 300000},
            {"category": "Generators", "count": 1, "fmv_mid_per_unit": 200000, "fmv_mid_subtotal": 200000},
        ],
        "totals": {"count": 3, "fmv_mid": 500000},
    }
    result = generate_onepager(data, user_message="Portfolio", client="Big Co", synthesis=synthesis)
    assert isinstance(result, bytes)
    text = _doc_text(data, user_message="Portfolio", client="Big Co", synthesis=synthesis)
    assert "Compressors" in text
    assert "Generators" in text


# ── edge cases ───────────────────────────────────────────────

def test_missing_valuation_defaults():
    data = {}
    result = generate_onepager(data, user_message="Unknown item")
    assert isinstance(result, bytes)
    assert result[:2] == b'PK'


def test_empty_client():
    data = {"valuation": {"fmv_mid": 10000, "currency": "CAD"}}
    result = generate_onepager(data, user_message="test", client="")
    assert isinstance(result, bytes)
